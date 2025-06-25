import os
import json
import uuid
from typing import List, Dict, Any, Optional
from pathlib import Path
from datetime import datetime

# Document processing imports
import fitz  # PyMuPDF for PDFs
from docx import Document  # python-docx for Word documents
import pandas as pd  # for CSV/Excel files
from sentence_transformers import SentenceTransformer
from qdrant_client import QdrantClient
from qdrant_client.http.models import (
    Distance, VectorParams, PointStruct, Filter,
    FieldCondition, MatchValue, UpdateStatus
)
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

class QdrantManager:
    """
    Simple Qdrant Vector Database Manager
    Handles document processing and basic CRUD operations
    """

    def __init__(self,
                 collection_name: str = None,
                 embedding_model: str = "all-MiniLM-L6-v2",
                 qdrant_url: str = None,
                 qdrant_api_key: str = None):
        """
        Initialize Qdrant Manager

        Args:
            collection_name: Name of Qdrant collection
            embedding_model: Sentence transformer model name
            qdrant_url: Qdrant server URL (if None, uses env variable)
            qdrant_api_key: Qdrant API key (if None, uses env variable)
        """
        self.collection_name = collection_name or f"documents_{int(datetime.now().timestamp())}"
        self.embedding_model = SentenceTransformer(embedding_model)
        self.embedding_dim = self.embedding_model.get_sentence_embedding_dimension()
        qdrant_url="https://488362f3-c049-424d-8c30-f08b860e85a4.eu-west-1-0.aws.cloud.qdrant.io"
        qdrant_api_key="eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJhY2Nlc3MiOiJtIn0.OxeUur4QZIrKbs4l-Et-r6x62pIYDcXOhAsJNGTtza8"
        # Qdrant connection details
        self.qdrant_url = qdrant_url or os.getenv("QDRANT_URL", "http://localhost:6333").rstrip("/")
        if self.qdrant_url.startswith("https://") and self.qdrant_url.endswith(":6333"):
            self.qdrant_url = self.qdrant_url.replace(":6333", "")

        self.qdrant_api_key = qdrant_api_key or os.getenv("QDRANT_API_KEY")

        # Initialize connection
        self.client = None
        self._setup_connection()
        self._setup_collection()

    def _setup_connection(self):
        """Initialize Qdrant client connection"""
        try:
            print(f"Connecting to Qdrant at: {self.qdrant_url}")

            if self.qdrant_api_key:
                self.client = QdrantClient(
                    url=self.qdrant_url,
                    api_key=self.qdrant_api_key
                )
            else:
                self.client = QdrantClient(url=self.qdrant_url)

            # Test connection
            collections = self.client.get_collections()
            print(f"Connected successfully. Found {len(collections.collections)} collections")

        except Exception as e:
            print(f"Failed to connect to Qdrant: {e}")
            raise

    def _setup_collection(self):
        """Create or verify collection setup"""
        try:
            # Check if collection exists
            try:
                collection_info = self.client.get_collection(self.collection_name)
                print(f"Using existing collection: {self.collection_name}")
                return
            except:
                pass  # Collection doesn't exist, create it

            # Create new collection
            print(f"Creating new collection: {self.collection_name}")
            self.client.create_collection(
                collection_name=self.collection_name,
                vectors_config=VectorParams(
                    size=self.embedding_dim,
                    distance=Distance.COSINE
                )
            )
            print(f"Collection '{self.collection_name}' created successfully")

        except Exception as e:
            print(f"Failed to setup collection: {e}")
            raise

    def _split_text_into_chunks(self, text: str, chunk_size: int = 500) -> List[str]:
        """Split text into smaller chunks"""
        words = text.split()
        chunks = []
        current_chunk = []
        current_length = 0

        for word in words:
            current_chunk.append(word)
            current_length += len(word) + 1

            if current_length >= chunk_size:
                chunk_text = " ".join(current_chunk)
                if len(chunk_text.strip()) > 50:
                    chunks.append(chunk_text)
                current_chunk = []
                current_length = 0

        # Add remaining words
        if current_chunk:
            chunk_text = " ".join(current_chunk)
            if len(chunk_text.strip()) > 50:
                chunks.append(chunk_text)

        return chunks

    def extract_text_from_pdf(self, pdf_path: str) -> List[Dict]:
        """Extract text from PDF and split into chunks"""
        try:
            doc = fitz.open(pdf_path)
            all_chunks = []

            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()

                if not text.strip():
                    continue

                chunks = self._split_text_into_chunks(text)
                for i, chunk in enumerate(chunks):
                    all_chunks.append({
                        'text': chunk,
                        'source': os.path.basename(pdf_path),
                        'file_type': 'pdf',
                        'page': page_num + 1,
                        'chunk_id': len(all_chunks)
                    })

            doc.close()
            return all_chunks

        except Exception as e:
            print(f"Error processing PDF {pdf_path}: {e}")
            return []

    def extract_text_from_docx(self, docx_path: str) -> List[Dict]:
        """Extract text from Word document and split into chunks"""
        try:
            doc = Document(docx_path)
            full_text = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))

            # Combine and chunk text
            combined_text = "\n".join(full_text)
            chunks = self._split_text_into_chunks(combined_text)

            result = []
            for i, chunk in enumerate(chunks):
                result.append({
                    'text': chunk,
                    'source': os.path.basename(docx_path),
                    'file_type': 'docx',
                    'chunk_id': i
                })

            return result

        except Exception as e:
            print(f"Error processing DOCX {docx_path}: {e}")
            return []

    def extract_text_from_txt(self, txt_path: str) -> List[Dict]:
        """Extract text from text file and split into chunks"""
        try:
            with open(txt_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()

            if not content.strip():
                return []

            chunks = self._split_text_into_chunks(content)
            result = []
            for i, chunk in enumerate(chunks):
                result.append({
                    'text': chunk,
                    'source': os.path.basename(txt_path),
                    'file_type': 'txt',
                    'chunk_id': i
                })

            return result

        except Exception as e:
            print(f"Error processing TXT {txt_path}: {e}")
            return []

    def extract_text_from_csv(self, csv_path: str) -> List[Dict]:
        """Extract text from CSV file"""
        try:
            df = pd.read_csv(csv_path)
            chunks = []

            for idx, row in df.iterrows():
                row_text = " | ".join([f"{col}: {str(val)}" for col, val in row.items() if pd.notna(val)])

                if len(row_text.strip()) > 20:
                    chunks.append({
                        'text': row_text,
                        'source': os.path.basename(csv_path),
                        'file_type': 'csv',
                        'row_number': idx + 1,
                        'chunk_id': len(chunks)
                    })

            return chunks

        except Exception as e:
            print(f"Error processing CSV {csv_path}: {e}")
            return []

    def extract_text_from_json(self, json_path: str) -> List[Dict]:
        """Extract text from JSON file"""
        try:
            with open(json_path, 'r', encoding='utf-8') as file:
                data = json.load(file)

            chunks = []

            def extract_from_obj(obj, path=""):
                if isinstance(obj, dict):
                    for key, value in obj.items():
                        new_path = f"{path}.{key}" if path else key
                        extract_from_obj(value, new_path)
                elif isinstance(obj, list):
                    for i, item in enumerate(obj):
                        new_path = f"{path}[{i}]"
                        extract_from_obj(item, new_path)
                else:
                    text = f"{path}: {str(obj)}"
                    if len(text.strip()) > 10:
                        chunks.append({
                            'text': text,
                            'source': os.path.basename(json_path),
                            'file_type': 'json',
                            'json_path': path,
                            'chunk_id': len(chunks)
                        })

            extract_from_obj(data)
            return chunks

        except Exception as e:
            print(f"Error processing JSON {json_path}: {e}")
            return []

    def process_file(self, file_path: str) -> List[Dict]:
        """Process a file based on its extension"""
        file_path = Path(file_path)
        extension = file_path.suffix.lower()

        print(f"Processing {extension} file: {file_path.name}")

        if extension == '.pdf':
            return self.extract_text_from_pdf(str(file_path))
        elif extension == '.docx':
            return self.extract_text_from_docx(str(file_path))
        elif extension == '.txt':
            return self.extract_text_from_txt(str(file_path))
        elif extension == '.csv':
            return self.extract_text_from_csv(str(file_path))
        elif extension == '.json':
            return self.extract_text_from_json(str(file_path))
        else:
            print(f"Unsupported file type: {extension}")
            return []

    def store_documents(self, chunks: List[Dict], batch_size: int = 30) -> bool:
        """Store document chunks in Qdrant"""
        if not chunks:
            print("No chunks to store")
            return False

        try:
            # Generate embeddings
            print("Generating embeddings...")
            texts = [chunk['text'] for chunk in chunks]
            embeddings = self.embedding_model.encode(texts, show_progress_bar=True)

            # Prepare points for insertion
            points = []
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                point = PointStruct(
                    id=str(uuid.uuid4()),
                    vector=embedding.tolist(),
                    payload={**chunk, 'stored_at': datetime.now().isoformat()}
                )
                points.append(point)

            # Insert in batches
            print(f"Storing {len(points)} points...")
            for i in range(0, len(points), batch_size):
                batch_points = points[i:i + batch_size]
                result = self.client.upsert(
                    collection_name=self.collection_name,
                    points=batch_points,
                    wait=True
                )

                if result.status != UpdateStatus.COMPLETED:
                    print(f"Failed to store batch {i//batch_size + 1}")
                    return False

            print(f"Successfully stored all {len(points)} documents")
            return True

        except Exception as e:
            print(f"Storage operation failed: {e}")
            return False

    def search(self,
              query: str,
              limit: int = 5,
              score_threshold: float = None,
              filter_conditions: Dict = None) -> List[Dict]:
        """Search for similar documents"""
        try:
            # Generate query embedding
            query_vector = self.embedding_model.encode([query])[0]

            # Prepare filter if provided
            search_filter = None
            if filter_conditions:
                search_filter = Filter(
                    must=[
                        FieldCondition(
                            key=key,
                            match=MatchValue(value=value)
                        ) for key, value in filter_conditions.items()
                    ]
                )

            # Perform search
            search_result = self.client.search(
                collection_name=self.collection_name,
                query_vector=query_vector.tolist(),
                query_filter=search_filter,
                limit=limit,
                score_threshold=score_threshold,
                with_payload=True
            )

            # Format results
            results = []
            for hit in search_result:
                results.append({
                    'id': hit.id,
                    'score': hit.score,
                    'text': hit.payload.get('text', ''),
                    'source': hit.payload.get('source', ''),
                    'file_type': hit.payload.get('file_type', ''),
                    'metadata': hit.payload
                })

            return results

        except Exception as e:
            print(f"Search failed: {e}")
            return []

def main():
    print("Qdrant Manager - Interactive Mode")
    print("=" * 40)

    collection_name = "docs"
    manager = QdrantManager(collection_name=collection_name)
    
    file_path = "/content/DelayLossThroughput_4.pdf"
    if file_path and os.path.exists(file_path):
        chunks = manager.process_file(file_path)
        if chunks:
            print(f"Extracted {len(chunks)} chunks from the file.")
            manager.store_documents(chunks)
        else:
            print("No chunks extracted from file.")
    else:
        print("File not found or invalid path.")

    query = input("Enter your search query: ").strip()
    if query:
        limit = 5
        results = manager.search(query, limit=limit)
        # ... rest of the search display code

        if results:
                print(f"\nFound {len(results)} results:")
                print("-" * 60)
                for i, result in enumerate(results, 1):
                    print(f"{i}. ID: {result['id']}")
                    print(f"   Score: {result['score']:.3f}")
                    print(f"   Source: {result['source']} ({result['file_type']})")
                    print(f"   Text: {result['text'][:200]}...")
                    print("-" * 60)

                    # Store results for potential updates/deletes
                    globals()['last_search_results'] = results
        else:
                    print("No results found.")
    else:
        print("Please enter a search query.")
if __name__ == "__main__":
    main()